"""
Utilities to ingest PDF/DOCX legal documents while preserving text, structure, and images.
"""

from __future__ import annotations

import re
import os
from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO, Iterable, List, Optional, Union
from io import BytesIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image as PILImage
try:
    import pytesseract

    OCR_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None
    OCR_AVAILABLE = False

# Support for .doc files (Word 97-2003)
# We'll convert .doc to .docx using LibreOffice or use python-docx2txt
try:
    import subprocess
    SUBPROCESS_AVAILABLE = True
except ImportError:
    SUBPROCESS_AVAILABLE = False


@dataclass
class SectionChunk:
    """Structured chunk extracted from a legal document."""

    level: str
    code: str
    title: str
    content: str
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    is_ocr: bool = False
    metadata: Optional[dict] = None


@dataclass
class ExtractedImage:
    """Image extracted from the source document."""

    data: bytes
    extension: str
    content_type: str
    page_number: Optional[int] = None
    description: str = ""
    width: Optional[int] = None
    height: Optional[int] = None


@dataclass
class ExtractedDocument:
    """Return value when parsing one document."""

    text: str
    page_count: int
    sections: List[SectionChunk]
    images: List[ExtractedImage]
    ocr_text: Optional[str] = None


SECTION_REGEX = re.compile(
    r"^(Chương\s+[IVXLC\d]+|Mục\s+[IVXLC\d]+|Điều\s+\d+[\w]*)",
    re.IGNORECASE,
)


def _detect_level(header: str) -> str:
    header_lower = header.lower()
    if header_lower.startswith("chương"):
        return "chapter"
    if header_lower.startswith("mục"):
        return "section"
    if header_lower.startswith("điều"):
        return "article"
    return "other"


def _split_sections(paragraphs: Iterable[str], *, is_ocr: bool = False) -> List[SectionChunk]:
    sections: List[SectionChunk] = []
    current: Optional[SectionChunk] = None

    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue

        match = SECTION_REGEX.match(paragraph)
        if match:
            header = match.group(0)
            rest = paragraph[len(header) :].strip()
            level = _detect_level(header)
            current = SectionChunk(
                level=level,
                code=header,
                title=rest,
                content=paragraph,
                is_ocr=is_ocr,
            )
            sections.append(current)
        elif current:
            current.content += "\n" + paragraph
        else:
            current = SectionChunk(
                level="other",
                code="Lời mở đầu",
                title="",
                content=paragraph,
                is_ocr=is_ocr,
            )
            sections.append(current)

    return sections


def _extract_docx_images(doc: DocxDocument) -> List[ExtractedImage]:
    images: List[ExtractedImage] = []
    rels = doc.part._rels.values()
    for rel in rels:
        if "image" not in rel.reltype:
            continue
        part = rel.target_part
        data = part.blob
        # Determine extension and metadata
        partname = Path(part.partname)
        ext = partname.suffix.lstrip(".") or "bin"
        content_type = getattr(part, "content_type", "application/octet-stream")
        width = None
        height = None
        try:
            with PILImage.open(BytesIO(data)) as pil_img:
                width, height = pil_img.size
        except Exception:
            pass
        images.append(
            ExtractedImage(
                data=data,
                extension=ext,
                content_type=content_type,
                page_number=None,
                width=width,
                height=height,
            )
        )
    return images


def extract_from_docx(path: Optional[Path] = None, data: Optional[bytes] = None) -> ExtractedDocument:
    """Parse DOCX file (path or bytes), keeping paragraphs in order and capturing embedded images."""
    if path is None and data is None:
        raise ValueError("DOCX extraction requires path or bytes.")
    if data is not None:
        doc = DocxDocument(BytesIO(data))
    else:
        doc = DocxDocument(path)
    paragraphs = [para.text for para in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    sections = _split_sections(paragraphs, is_ocr=False)
    images = _extract_docx_images(doc)
    # DOCX has no fixed page count; approximate by paragraphs length
    sections = _apply_chunk_strategy(sections, full_text)
    return ExtractedDocument(
        text=full_text,
        page_count=len(doc.paragraphs) or 1,
        sections=sections,
        images=images,
        ocr_text=None,
    )


def _pixmap_to_pil(pix: fitz.Pixmap) -> PILImage.Image:
    mode = "RGB"
    if pix.n == 1:
        mode = "L"
    elif pix.n == 4:
        mode = "RGBA"
    return PILImage.frombytes(mode, [pix.width, pix.height], pix.samples)


def _perform_ocr_on_page(page: fitz.Page) -> str:
    if not OCR_AVAILABLE:
        return ""
    try:
        zoom = os.getenv("OCR_PDF_ZOOM", "2.0")
        try:
            zoom_val = float(zoom)
        except ValueError:
            zoom_val = 2.0
        matrix = fitz.Matrix(zoom_val, zoom_val)
        pix = page.get_pixmap(matrix=matrix)
        pil_img = _pixmap_to_pil(pix)
        langs = os.getenv("OCR_LANGS", "vie+eng")
        text = pytesseract.image_to_string(pil_img, lang=langs)
        return text.strip()
    except Exception:
        return ""


def _extract_pdf_images(pdf: fitz.Document) -> List[ExtractedImage]:
    images: List[ExtractedImage] = []
    for page_index in range(pdf.page_count):
        page = pdf.load_page(page_index)
        for image in page.get_images(full=True):
            xref = image[0]
            try:
                pix = fitz.Pixmap(pdf, xref)
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_bytes = pix.tobytes("png")
                images.append(
                    ExtractedImage(
                        data=img_bytes,
                        extension="png",
                        content_type="image/png",
                        page_number=page_index + 1,
                        width=pix.width,
                        height=pix.height,
                    )
                )
                if pix.alpha and pix.n > 4:
                    pix = None
            except Exception:
                continue
    return images


def extract_from_doc(path: Optional[Path] = None, data: Optional[bytes] = None) -> ExtractedDocument:
    """
    Parse .doc file (Word 97-2003 format).
    Converts .doc to .docx using LibreOffice if available, then processes as .docx.
    Otherwise, extracts text using basic methods.
    """
    if path is None and data is None:
        raise ValueError("DOC extraction requires path or bytes.")
    
    import tempfile
    import shutil
    
    # If we have data, save to temp file
    if data is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
            tmp.write(data)
            doc_path = Path(tmp.name)
            temp_created = True
    else:
        doc_path = Path(path)
        temp_created = False
    
    try:
        # Try to convert .doc to .docx using LibreOffice
        if SUBPROCESS_AVAILABLE:
            try:
                # Check if LibreOffice is available
                result = subprocess.run(
                    ['which', 'libreoffice'] if os.name != 'nt' else ['where', 'libreoffice'],
                    capture_output=True,
                    text=True
                )
                if result.returncode == 0 or shutil.which('libreoffice') or shutil.which('soffice'):
                    # Convert .doc to .docx
                    with tempfile.TemporaryDirectory() as tmpdir:
                        output_dir = Path(tmpdir)
                        # Use soffice (LibreOffice) or libreoffice command
                        cmd = shutil.which('soffice') or shutil.which('libreoffice')
                        if cmd:
                            subprocess.run(
                                [cmd, '--headless', '--convert-to', 'docx', '--outdir', str(output_dir), str(doc_path)],
                                check=True,
                                capture_output=True,
                                timeout=30
                            )
                            # Find the converted file
                            converted_file = output_dir / (doc_path.stem + '.docx')
                            if converted_file.exists():
                                # Process as .docx
                                return extract_from_docx(path=converted_file)
            except (subprocess.SubprocessError, FileNotFoundError, TimeoutError):
                pass  # Fall through to basic text extraction
        
        # Fallback: Basic text extraction using python-docx (won't work for .doc)
        # Or try to read as plain text
        try:
            # Try to read as text (basic fallback)
            with open(doc_path, 'rb') as f:
                # Skip binary header, try to extract readable text
                content = f.read()
                # Very basic: try to extract text between null bytes or readable ranges
                # This is a last resort and won't work well
                text_parts = []
                current_text = ""
                for byte in content:
                    if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII
                        current_text += chr(byte)
                    else:
                        if len(current_text) > 10:
                            text_parts.append(current_text)
                        current_text = ""
                if current_text:
                    text_parts.append(current_text)
                
                full_text = "\n".join(text_parts)
                if len(full_text) > 100:  # If we got reasonable text
                    paragraphs = [p.strip() for p in full_text.split('\n') if p.strip()]
                    sections = _split_sections(paragraphs, is_ocr=False)
                    sections = _apply_chunk_strategy(sections, full_text)
                    return ExtractedDocument(
                        text=full_text,
                        page_count=len(paragraphs) or 1,
                        sections=sections,
                        images=[],
                        ocr_text=None,
                    )
        except Exception:
            pass
        
        # If all else fails, raise helpful error
        raise ValueError(
            "File type .doc (Word 97-2003) is not fully supported. "
            "Please convert the file to .docx format using Microsoft Word or LibreOffice, "
            "or install LibreOffice command-line tools for automatic conversion."
        )
    finally:
        if temp_created and doc_path.exists():
            os.unlink(doc_path)


def extract_from_pdf(path: Optional[Path] = None, data: Optional[bytes] = None) -> ExtractedDocument:
    """Parse PDF file using PyMuPDF (path or bytes) and capture page text + images."""
    if path is None and data is None:
        raise ValueError("PDF extraction requires path or bytes.")
    if data is not None:
        pdf = fitz.open(stream=data, filetype="pdf")
    else:
        pdf = fitz.open(path)

    fragments: List[str] = []
    ocr_fragments: List[str] = []
    sections: List[SectionChunk] = []
    current: Optional[SectionChunk] = None

    for page_index in range(pdf.page_count):
        page = pdf.load_page(page_index)
        page_text = page.get_text("text").strip()
        page_is_ocr = False
        if not page_text:
            ocr_text = _perform_ocr_on_page(page)
            if ocr_text:
                page_text = ocr_text
                page_is_ocr = True
                ocr_fragments.append(ocr_text)
        fragments.append(page_text)

        for paragraph in page_text.splitlines():
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            match = SECTION_REGEX.match(paragraph)
            if match:
                header = match.group(0)
                rest = paragraph[len(header) :].strip()
                level = _detect_level(header)
                current = SectionChunk(
                    level=level,
                    code=header,
                    title=rest,
                    content=paragraph,
                    page_start=page_index + 1,
                    page_end=page_index + 1,
                    is_ocr=page_is_ocr,
                )
                sections.append(current)
            elif current:
                current.content += "\n" + paragraph
                current.page_end = page_index + 1
                current.is_ocr = current.is_ocr or page_is_ocr
            else:
                current = SectionChunk(
                    level="other",
                    code="Trang đầu",
                    title="",
                    content=paragraph,
                    page_start=page_index + 1,
                    page_end=page_index + 1,
                    is_ocr=page_is_ocr,
                )
                sections.append(current)

    images = _extract_pdf_images(pdf)
    full_text = "\n".join(fragments)
    ocr_text = "\n".join(ocr_fragments) if ocr_fragments else None
    sections = _apply_chunk_strategy(sections, full_text)
    return ExtractedDocument(
        text=full_text,
        page_count=pdf.page_count,
        sections=sections,
        images=images,
        ocr_text=ocr_text,
    )


def _generate_semantic_chunks(text: str, chunk_size: int, overlap: int) -> List[SectionChunk]:
    if chunk_size <= 0:
        return []
    overlap = max(0, min(overlap, chunk_size - 1))
    chunks: List[SectionChunk] = []
    length = len(text)
    start = 0
    idx = 1
    while start < length:
        end = min(length, start + chunk_size)
        chunk_content = text[start:end].strip()
        if chunk_content:
            chunks.append(
                SectionChunk(
                    level="chunk",
                    code=f"Chunk {idx}",
                    title="",
                    content=chunk_content,
                    metadata={"chunk_strategy": "semantic"},
                )
            )
            idx += 1
        if end >= length:
            break
        start = max(0, end - overlap)
    return chunks


def _apply_chunk_strategy(sections: List[SectionChunk], full_text: str) -> List[SectionChunk]:
    strategy = os.getenv("LEGAL_CHUNK_STRATEGY", "structure").lower()
    if strategy != "hybrid":
        return sections
    try:
        chunk_size = int(os.getenv("LEGAL_CHUNK_SIZE", "1200"))
    except ValueError:
        chunk_size = 1200
    try:
        overlap = int(os.getenv("LEGAL_CHUNK_OVERLAP", "200"))
    except ValueError:
        overlap = 200
    new_sections = list(sections)
    new_sections.extend(_generate_semantic_chunks(full_text, chunk_size, overlap))
    return new_sections


SourceType = Union[str, Path, BinaryIO]


def load_legal_document(source: SourceType, filename: Optional[str] = None) -> ExtractedDocument:
    """
    Dispatch helper depending on file type.

    Args:
        source: path or binary handle.
        filename: optional original filename (needed when source is a stream).

    Raises:
        ValueError: if extension unsupported.
    """
    path_obj: Optional[Path] = None
    data: Optional[bytes] = None

    if isinstance(source, (str, Path)):
        path_obj = Path(source)
        suffix = path_obj.suffix.lower()
    else:
        data = source.read()
        if hasattr(source, "seek"):
            source.seek(0)
        suffix = Path(filename or "").suffix.lower()

    if suffix == ".docx":
        return extract_from_docx(path=path_obj, data=data)
    if suffix == ".doc":
        return extract_from_doc(path=path_obj, data=data)
    if suffix == ".pdf":
        return extract_from_pdf(path=path_obj, data=data)
    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

