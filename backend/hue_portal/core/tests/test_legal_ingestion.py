import os
import shutil
import tempfile
from io import BytesIO

from django.test import TestCase, override_settings
from django.core.files.uploadedfile import SimpleUploadedFile
from PIL import Image as PILImage
from docx import Document

from hue_portal.core.services import ingest_uploaded_document, enqueue_ingestion_job
from hue_portal.core.models import LegalDocument, IngestionJob


class LegalIngestionServiceTests(TestCase):
    def setUp(self):
        self.media_dir = tempfile.mkdtemp(prefix="legal-media-")
        self.override = override_settings(MEDIA_ROOT=self.media_dir)
        self.override.enable()

    def tearDown(self):
        self.override.disable()
        shutil.rmtree(self.media_dir, ignore_errors=True)

    def _make_docx_with_image(self) -> bytes:
        document = Document()
        document.add_paragraph("Điều 1. Quy định chung")
        document.add_paragraph("Nội dung điều 1 được ghi rõ ràng.")

        fd, image_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        try:
            pil_image = PILImage.new("RGB", (32, 32), color="red")
            pil_image.save(image_path)
            document.add_picture(image_path)
        finally:
            os.remove(image_path)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _make_docx_with_header(self, header: str, body: str) -> bytes:
        document = Document()
        document.add_paragraph(header)
        for line in body.split("\n"):
            document.add_paragraph(line)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def test_ingest_docx_extracts_sections_and_images(self):
        docx_bytes = self._make_docx_with_image()
        metadata = {
            "code": "TEST-DOC-1",
            "title": "Tài liệu thử nghiệm",
            "doc_type": "circular",
            "summary": "Tài liệu test",
            "issued_by": "Test Unit",
            "issued_at": "2025-11-18",
            "source_url": "",
            "metadata": {"tags": ["demo"]},
        }

        result = ingest_uploaded_document(
            file_obj=BytesIO(docx_bytes),
            filename="test.docx",
            metadata=metadata,
        )

        self.assertGreaterEqual(result.sections_count, 1)
        self.assertEqual(result.images_count, 1)
        self.assertTrue(result.document.raw_text.startswith("Điều 1"))
        self.assertTrue(result.document.file_checksum)
        self.assertEqual(result.document.raw_text_ocr, "")
        self.assertTrue(result.document.uploaded_file.name)
        self.assertTrue(result.document.images.exists())

        stored_doc = LegalDocument.objects.get(code="TEST-DOC-1")
        self.assertGreaterEqual(stored_doc.sections.count(), 1)
        self.assertEqual(stored_doc.sections.filter(is_ocr=True).count(), 0)

    def test_enqueue_ingestion_job_runs_when_eager(self):
        docx_bytes = self._make_docx_with_image()
        upload = SimpleUploadedFile("test.docx", docx_bytes, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        metadata = {
            "code": "TEST-DOC-QUEUE",
            "title": "Hàng đợi",
            "doc_type": "decision",
        }

        job = enqueue_ingestion_job(file_obj=upload, filename=upload.name, metadata=metadata)
        job.refresh_from_db()

        self.assertEqual(job.status, IngestionJob.STATUS_COMPLETED)
        self.assertIsNotNone(job.document)
        self.assertEqual(job.stats.get("sections"), job.document.sections.count())

    def test_auto_metadata_and_deduplication(self):
        header = "QUYẾT ĐỊNH CỦA BỘ CÔNG AN\nNgày 01/02/2024"
        docx_bytes = self._make_docx_with_header(header, "Nội dung quyết định ...")
        metadata = {
            "code": "AUTO-META",
            "title": "",
            "doc_type": "other",
            "issued_by": "",
            "issued_at": "",
        }
        result = ingest_uploaded_document(
            file_obj=BytesIO(docx_bytes),
            filename="auto.docx",
            metadata=metadata,
        )
        stored_doc = LegalDocument.objects.get(code="AUTO-META")
        self.assertEqual(stored_doc.doc_type, "decision")
        self.assertIsNotNone(stored_doc.issued_at)
        self.assertIn("Bộ Công An", stored_doc.issued_by.title())
        self.assertTrue(result.document.content_checksum)

        metadata_dup = {
            "code": "AUTO-META-2",
            "title": "",
            "doc_type": "other",
        }
        with self.assertRaises(ValueError):
            ingest_uploaded_document(
                file_obj=BytesIO(docx_bytes),
                filename="auto-copy.docx",
                metadata=metadata_dup,
            )

